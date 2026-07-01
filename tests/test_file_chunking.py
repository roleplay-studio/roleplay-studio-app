"""Tests for file parsing and chunking in knowledge base."""

from __future__ import annotations

import io

import pytest

from app.application.dto import KnowledgeEntryDTO
from app.application.services.file_parser import FileParser
from app.application.services.text_chunker import TextChunker

# ── Tests: FileParser ─────────────────────────────────────────────────


class TestFileParser:
    def test_parse_txt_file(self):
        """Parse plain text file."""
        content = b"Hello world\nThis is a test file.\nThird line."
        result = FileParser.parse(content, "test.txt", "text/plain")
        assert result == "Hello world\nThis is a test file.\nThird line."

    def test_parse_md_file(self):
        """Parse markdown file."""
        content = b"# Title\n\nSome **bold** text.\n\n## Section\n\nContent here."
        result = FileParser.parse(content, "test.md", "text/markdown")
        assert "# Title" in result
        assert "Some **bold** text" in result

    def test_parse_txt_by_extension(self):
        """Detect file type by extension when mime is generic."""
        content = b"Plain text content"
        result = FileParser.parse(content, "notes.txt", "application/octet-stream")
        assert result == "Plain text content"

    def test_parse_empty_file(self):
        """Empty file returns empty string."""
        result = FileParser.parse(b"", "empty.txt", "text/plain")
        assert result == ""

    def test_parse_unsupported_format_raises(self):
        """Unsupported file format raises ValueError."""
        with pytest.raises(ValueError, match="Unsupported"):
            FileParser.parse(b"data", "file.xyz", "application/xyz")

    def test_parse_docx_file(self):
        """Parse DOCX file using python-docx."""
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        result = FileParser.parse(
            content,
            "test.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert "Test Document" in result
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_parse_pdf_file(self):
        """Parse PDF file using pypdf."""
        pytest.importorskip("pypdf")
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)

        buf = io.BytesIO()
        writer.write(buf)
        content = buf.getvalue()

        # PDF with blank page should parse without error
        result = FileParser.parse(content, "test.pdf", "application/pdf")
        assert isinstance(result, str)


# ── Tests: TextChunker ────────────────────────────────────────────────


class TestTextChunker:
    def test_chunk_short_text_returns_single_chunk(self):
        """Text shorter than chunk_size should return as single chunk."""
        chunker = TextChunker(chunk_size=1000, chunk_overlap=100)
        chunks = chunker.chunk("Short text.")
        assert len(chunks) == 1
        assert chunks[0] == "Short text."

    def test_chunk_long_text_splits(self):
        """Long text should be split into multiple chunks."""
        text = "Word " * 500  # 2500 chars
        chunker = TextChunker(chunk_size=500, chunk_overlap=50)
        chunks = chunker.chunk(text)
        assert len(chunks) > 1
        # All chunks should be non-empty
        for chunk in chunks:
            assert len(chunk.strip()) > 0

    def test_chunk_preserves_all_content(self):
        """All original content should appear in chunks (possibly with overlap)."""
        text = "Sentence one. Sentence two. Sentence three. Sentence four. Sentence five."
        chunker = TextChunker(chunk_size=50, chunk_overlap=10)
        chunks = chunker.chunk(text)
        combined = " ".join(chunks)
        assert "Sentence one" in combined
        assert "Sentence five" in combined

    def test_chunk_empty_text_returns_empty(self):
        """Empty text returns empty list."""
        chunker = TextChunker()
        chunks = chunker.chunk("")
        assert chunks == []

    def test_chunk_with_metadata(self):
        """chunk_with_metadata returns list of (text, metadata) tuples."""
        chunker = TextChunker(chunk_size=1000, chunk_overlap=100)
        result = chunker.chunk_with_metadata(
            "Some text content.",
            source_type="file",
            file_name="test.txt",
        )
        assert len(result) == 1
        text, meta = result[0]
        assert text == "Some text content."
        assert meta["source_type"] == "file"
        assert meta["file_name"] == "test.txt"
        assert meta["chunk_index"] == 0

    def test_chunk_with_metadata_multiple(self):
        """Multiple chunks get correct chunk_index."""
        text = "Word " * 500
        chunker = TextChunker(chunk_size=500, chunk_overlap=50)
        result = chunker.chunk_with_metadata(
            text,
            source_type="file",
            file_name="big.txt",
        )
        assert len(result) > 1
        for i, (_, meta) in enumerate(result):
            assert meta["chunk_index"] == i
            assert meta["file_name"] == "big.txt"


# ── Tests: KnowledgeEntryDTO extended ─────────────────────────────────


class TestKnowledgeEntryDTOExtended:
    def test_dto_has_source_type_field(self):
        """KnowledgeEntryDTO should have source_type field."""
        entry = KnowledgeEntryDTO(id="1", content="text", source_type="manual")
        assert entry.source_type == "manual"

    def test_dto_source_type_default_manual(self):
        """Default source_type should be 'manual'."""
        entry = KnowledgeEntryDTO(id="1", content="text")
        assert entry.source_type == "manual"

    def test_dto_has_file_name_field(self):
        """KnowledgeEntryDTO should have optional file_name field."""
        entry = KnowledgeEntryDTO(id="1", content="text", source_type="file", file_name="test.txt")
        assert entry.file_name == "test.txt"

    def test_dto_has_chunk_count_field(self):
        """KnowledgeEntryDTO should have chunk_count field."""
        entry = KnowledgeEntryDTO(id="1", content="text", chunk_count=5)
        assert entry.chunk_count == 5


# ── Tests: File upload command ────────────────────────────────────────


class TestAddKnowledgeFileCommand:
    def test_command_has_file_fields(self):
        """AddKnowledgeFileCommand should have file-related fields."""
        from app.application.dto import AddKnowledgeFileCommand

        cmd = AddKnowledgeFileCommand(
            bot_id=1,
            file_content=b"Hello",
            file_name="test.txt",
            mime_type="text/plain",
        )
        assert cmd.bot_id == 1
        assert cmd.file_name == "test.txt"
